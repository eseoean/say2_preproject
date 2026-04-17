from __future__ import annotations

from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/skku_aws2_18/pre_project/say2_preproject")
HTML_PATH = ROOT / "docs" / "preproject_submission_report_20260413_source.html"
DOCX_PATH = ROOT / "docs" / "preproject_submission_report_20260413.docx"


def extract_text(node: Tag) -> str:
    return " ".join(node.stripped_strings)


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)


def shade_header(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F6FB")
    tc_pr.append(shd)


def add_paragraph(document: Document, text: str, *, align=None, bold=False, size=11, italic=False) -> None:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(16)
    elif level == 3:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11.5)


def add_list(document: Document, node: Tag, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for item in node.find_all("li", recursive=False):
        paragraph = document.add_paragraph(style=style)
        run = paragraph.add_run(extract_text(item))
        run.font.size = Pt(11)


def add_table(document: Document, node: Tag) -> None:
    rows = node.find_all("tr")
    if not rows:
        return

    first_row_cells = rows[0].find_all(["th", "td"], recursive=False)
    table = document.add_table(rows=0, cols=len(first_row_cells))
    table.style = "Table Grid"

    for row_idx, row in enumerate(rows):
        cells = row.find_all(["th", "td"], recursive=False)
        if not cells:
            continue
        table_row = table.add_row()
        for col_idx, cell in enumerate(cells):
            target = table_row.cells[col_idx]
            set_cell_text(target, extract_text(cell))
            if row_idx == 0:
                for paragraph in target.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                shade_header(target)


def add_figure(document: Document, node: Tag) -> None:
    image = node.find("img")
    caption = node.find(class_="caption")
    if image is None:
        return

    image_path = (HTML_PATH.parent / image["src"]).resolve()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.7))

    if caption is not None:
        add_paragraph(
            document,
            extract_text(caption),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            italic=True,
            size=10,
        )


def build_cover(document: Document, node: Tag) -> None:
    subtitle = node.find(class_="cover-subtitle")
    title = node.find("h1")
    meta_lines = node.find(class_="cover-meta")

    if subtitle is not None:
        add_paragraph(document, extract_text(subtitle), align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    if title is not None:
        add_heading(document, title.get_text("\n", strip=True), 1)
    if meta_lines is not None:
        document.add_paragraph()
        for line in meta_lines.find_all("p", recursive=False):
            add_paragraph(document, extract_text(line), align=WD_ALIGN_PARAGRAPH.CENTER, size=11)


def apply_default_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)

    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def build_docx() -> None:
    soup = BeautifulSoup(HTML_PATH.read_text(encoding="utf-8"), "html.parser")
    body = soup.body
    if body is None:
        raise RuntimeError("HTML body not found")

    document = Document()
    apply_default_style(document)

    for child in body.children:
        if isinstance(child, NavigableString):
            continue
        if not isinstance(child, Tag):
            continue

        if child.name == "div" and "cover" in child.get("class", []):
            build_cover(document, child)
            continue
        if child.name == "div" and "page-break" in child.get("class", []):
            document.add_page_break()
            continue
        if child.name == "div" and "figure" in child.get("class", []):
            add_figure(document, child)
            continue
        if child.name == "h1":
            add_heading(document, child.get_text("\n", strip=True), 1)
            continue
        if child.name == "h2":
            add_heading(document, extract_text(child), 2)
            continue
        if child.name == "h3":
            add_heading(document, extract_text(child), 3)
            continue
        if child.name == "h4":
            add_heading(document, extract_text(child), 4)
            continue
        if child.name == "p":
            classes = child.get("class", [])
            size = 10 if "small" in classes else 11
            add_paragraph(document, extract_text(child), size=size)
            continue
        if child.name == "ul":
            add_list(document, child, ordered=False)
            continue
        if child.name == "ol":
            add_list(document, child, ordered=True)
            continue
        if child.name == "table":
            add_table(document, child)
            continue
        if child.name == "div" and "note" in child.get("class", []):
            add_paragraph(document, extract_text(child), italic=True, size=10.5)

    # Avoid an empty trailing section produced by the cover page logic.
    if document.sections[-1].start_type == WD_SECTION.NEW_PAGE:
        document.sections[-1].start_type = WD_SECTION.CONTINUOUS

    document.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
